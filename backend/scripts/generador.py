import os
import sys
import shutil
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docxtpl import DocxTemplate
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CONFIGURACIÓN VISUAL ---
ANCHO_COLUMNA_GRILLA = Inches(3.05) 
ANCHO_IMAGEN_GRILLA = Inches(3.0)
ALTO_IMAGEN_GRILLA  = Inches(2.25)
ANCHO_IMAGEN_SOLA = Inches(4.0)

# --- UTILIDADES DE ORDENAMIENTO ---
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

# --- UTILIDADES XML Y FORMATO ---
def aplicar_borde_imagen_xml(run):
    if not run.element.drawing_lst: return
    try:
        drawing = run.element.drawing_lst[0]
        graphic = None
        inline = drawing.find(qn('wp:inline'))
        if inline is not None: graphic = inline.find(qn('a:graphic'))
        else:
            anchor = drawing.find(qn('wp:anchor'))
            if anchor is not None: graphic = anchor.find(qn('a:graphic'))
        if graphic is None: return
        graphicData = graphic.find(qn('a:graphicData'))
        if graphicData is None: return
        pic = graphicData.find(qn('pic:pic'))
        if pic is None: return
        spPr = pic.find(qn('pic:spPr'))
        if spPr is None: return
        ln = OxmlElement('a:ln')
        ln.set('w', '31750') 
        solidFill = OxmlElement('a:solidFill')
        srgbClr = OxmlElement('a:srgbClr')
        srgbClr.set('val', '000000')
        solidFill.append(srgbClr)
        ln.append(solidFill)
        spPr.append(ln)
    except: pass

def set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m in ['top', 'bottom', 'start', 'end']:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), "0")
        node.set(qn("w:type"), "dxa")

def apagar_negrita_run(run):
    try:
        run.font.bold = False 
        run.font.color.rgb = RGBColor(0,0,0)
        r_rPr = run._element.get_or_add_rPr()
        for tag in ['w:b', 'w:bCs']:
            elements = r_rPr.findall(qn(tag))
            for elem in elements: r_rPr.remove(elem)
        rb = OxmlElement('w:b'); rb.set(qn('w:val'), '0'); r_rPr.append(rb)
        rbCs = OxmlElement('w:bCs'); rbCs.set(qn('w:val'), '0'); r_rPr.append(rbCs)
    except: pass

def limpiar_variables_en_parrafo(paragraph, valores_target):
    if not paragraph.text.strip(): return
    for run in paragraph.runs:
        texto_run = run.text
        if not texto_run.strip(): continue
        for val in valores_target:
            if val in texto_run or (len(texto_run) > 2 and texto_run in val):
                apagar_negrita_run(run)
                break 

def limpiar_formato_variables(doc, datos_raw):
    valores_target = []
    for k, v in datos_raw.items():
        val_str = str(v).strip()
        if val_str: valores_target.append(val_str)
    if not valores_target: return
    for para in doc.paragraphs: limpiar_variables_en_parrafo(para, valores_target)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs: limpiar_variables_en_parrafo(para, valores_target)

def aplicar_formato_comentario(parrafo):
    parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parrafo.paragraph_format.left_indent = Inches(0.5) 
    if parrafo.runs:
        for run in parrafo.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(10)
            apagar_negrita_run(run) 

def agregar_texto_comentario(doc, texto):
    if not texto: return
    p = doc.add_paragraph(texto)
    aplicar_formato_comentario(p)
    p.paragraph_format.space_after = Pt(12)

def leer_general_txt(ruta_txt):
    datos = {}
    comentarios = {}
    if not os.path.exists(ruta_txt): return {}, {}
    try:
        with open(ruta_txt, 'r', encoding='utf-8') as f: lines = f.readlines()
    except UnicodeDecodeError:
        with open(ruta_txt, 'r', encoding='latin-1') as f: lines = f.readlines()
    
    current_section = None
    current_val = []
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith("#"): continue
        if line.startswith("[") and line.endswith("]"):
            if current_section:
                if current_section == "COMENTARIOS_FOTOS":
                    for l in current_val:
                        if "//" in l:
                            k, v = l.split("//", 1)
                            comentarios[k.strip()] = v.strip()
                elif current_section != "DATOS":
                    datos[current_section] = "\n".join(current_val)
            
            current_section = line[1:-1].strip()
            current_val = []
        elif "=" in line and (current_section is None or current_section == "DATOS"):
            parts = line.split("=", 1)
            datos[parts[0].strip()] = parts[1].strip()
        elif current_section:
            current_val.append(line)
            
    if current_section:
        if current_section == "COMENTARIOS_FOTOS":
            for l in current_val:
                if "//" in l:
                    k, v = l.split("//", 1)
                    comentarios[k.strip()] = v.strip()
        elif current_section != "DATOS":
            datos[current_section] = "\n".join(current_val)
            
    return datos, comentarios

def agregar_foto_subtitulo(doc, ruta_img, subtitulo, comentario_top=None):
    if comentario_top:
        agregar_texto_comentario(doc, comentario_top)
    tabla = doc.add_table(rows=1, cols=1)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER 
    cell = tabla.cell(0,0)
    set_cell_margins_zero(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4) 
    run = p.add_run()
    run.add_picture(ruta_img, width=ANCHO_IMAGEN_SOLA) 
    aplicar_borde_imagen_xml(run)
    p_desc = doc.add_paragraph(subtitulo)
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.style.font.size = Pt(9)
    p_desc.style.font.bold = True
    p_desc.paragraph_format.space_after = Pt(12)

def descargar_grilla_fotos(doc, buffer_fotos):
    if not buffer_fotos: return
    for i in range(0, len(buffer_fotos), 2):
        par = buffer_fotos[i:i+2]
        tabla = doc.add_table(rows=1, cols=2)
        tabla.autofit = False 
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER 
        for col in tabla.columns: col.width = ANCHO_COLUMNA_GRILLA
        for j, (ruta_img, desc) in enumerate(par):
            cell = tabla.rows[0].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins_zero(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run()
            run.add_picture(ruta_img, width=ANCHO_IMAGEN_GRILLA, height=ALTO_IMAGEN_GRILLA)
            aplicar_borde_imagen_xml(run)
            if desc: 
                p_txt = cell.add_paragraph(desc)
                p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_txt.paragraph_format.space_before = Pt(4)
                p_txt.paragraph_format.space_after = Pt(4)
                run_txt = p_txt.runs[0] if p_txt.runs else p_txt.add_run(desc)
                run_txt.font.size = Pt(9)
        doc.add_paragraph("").paragraph_format.space_after = Pt(12)

def procesar_evidencias_universal(doc, carpeta_nodo, comentarios_dict):
    elementos = sorted([f for f in os.listdir(carpeta_nodo) if f != "GENERAL.txt" and f != "temp_render.docx" and not f.endswith(".docx") and not f.endswith(".pdf")], key=natural_sort_key)
    
    print(f"   [DIR] Procesando carpeta: {os.path.basename(carpeta_nodo)}")
    
    for nombre in elementos:
        ruta_completa = os.path.join(carpeta_nodo, nombre)
        if nombre.startswith("."): continue
        
        if "_" in nombre:
            parts = nombre.split("_", 1)
            titulo_visible = parts[1].strip()
            comentario = comentarios_dict.get(nombre, None)
            
            if os.path.isfile(ruta_completa) and nombre.lower().endswith(('.jpg','.jpeg','.png')):
                 titulo_visible = os.path.splitext(titulo_visible)[0]
                 agregar_foto_subtitulo(doc, ruta_completa, titulo_visible, comentario)
            elif os.path.isdir(ruta_completa):
                fotos_dentro = [f for f in os.listdir(ruta_completa) if f.lower().endswith(('.jpg','.jpeg','.png'))]
                if fotos_dentro:
                    img_path = os.path.join(ruta_completa, fotos_dentro[0])
                    agregar_foto_subtitulo(doc, img_path, titulo_visible, comentario)
        
        elif "-" in nombre:
            parts = nombre.split("-", 1)
            alineacion = WD_ALIGN_PARAGRAPH.LEFT 
            titulo_texto = parts[1].strip()
            
            if titulo_texto.startswith("--"): 
                alineacion = WD_ALIGN_PARAGRAPH.RIGHT
                titulo_texto = titulo_texto[2:].strip()
            elif titulo_texto.startswith("-"):
                alineacion = WD_ALIGN_PARAGRAPH.CENTER
                titulo_texto = titulo_texto[1:].strip()
            
            p_tit = doc.add_paragraph(titulo_texto)
            p_tit.alignment = alineacion
            run_tit = p_tit.runs[0] if p_tit.runs else p_tit.add_run(titulo_texto)
            run_tit.font.bold = True
            run_tit.font.size = Pt(10)
            p_tit.paragraph_format.space_after = Pt(12)
            
            comentario = comentarios_dict.get(nombre, None)
            if comentario:
                agregar_texto_comentario(doc, comentario)
            
            buffer_fotos = []
            if os.path.isdir(ruta_completa):
                fotos = sorted([f for f in os.listdir(ruta_completa) if f.lower().endswith(('.jpg','.jpeg','.png'))], key=natural_sort_key)
                for f in fotos:
                    f_path = os.path.join(ruta_completa, f)
                    filename_no_ext = os.path.splitext(f)[0]
                    caption = ""
                    if "_" in filename_no_ext:
                         parts = filename_no_ext.split("_", 1)
                         caption = parts[1].strip()
                    buffer_fotos.append((f_path, caption))
            elif os.path.isfile(ruta_completa) and nombre.lower().endswith(('.jpg','.jpeg','.png')):
                 filename_no_ext = os.path.splitext(nombre)[0]
                 caption = ""
                 if "_" in filename_no_ext:
                      parts = filename_no_ext.split("_", 1)
                      caption = parts[1].strip()
                 buffer_fotos.append((ruta_completa, caption))
            descargar_grilla_fotos(doc, buffer_fotos)

# --- PUNTO DE ENTRADA PARA NODE.JS ---
if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Error: Faltan argumentos (Ruta y Modo)")
        sys.exit(1)

    ruta_proyecto = sys.argv[1]
    modo = sys.argv[2] 
    
    ruta_base = os.path.dirname(os.path.abspath(__file__))
    ruta_plantilla = os.path.join(ruta_base, "../templates/plantilla_general.docx")

    print(f"--- INICIO PYTHON: {ruta_proyecto} [{modo}] ---")

    ruta_txt = os.path.join(ruta_proyecto, "GENERAL.txt")
    if not os.path.exists(ruta_txt):
        print("Error: No existe GENERAL.txt")
        sys.exit(1)

    datos, comentarios = leer_general_txt(ruta_txt)
    
    # === NUEVA LÓGICA DE NOMBRE DE ARCHIVO v8.3 ===
    
    # 1. Obtener valores limpios
    codigo_proy = datos.get('INFORME_NUM', '000').strip()
    codigo_cli = datos.get('CODIGO_CLIENTE', '').strip()
    empresa = datos.get('EMPRESA', 'Cliente').strip()
    
    distrito = datos.get('DISTRITO', '').strip()
    direccion = datos.get('DIRECCION', '').strip()
    
    # 2. Construir Ubicación (DIRECCION-DISTRITO)
    # Si no existen, fallback a 'LOCAL' antiguo
    if not distrito and not direccion:
        ubicacion = datos.get('LOCAL', 'Sede').strip()
    else:
        # Usamos guion para separar direccion de distrito
        partes_ubi = []
        if direccion: partes_ubi.append(direccion) # Primero Dirección
        if distrito: partes_ubi.append(distrito)   # Luego Distrito
        ubicacion = "-".join(partes_ubi)

    # 3. Función de limpieza
    def limpiar_nombre(s):
        s = str(s).replace("/", "-").replace("\\", "-").replace(":", "").replace("*", "").replace("?", "").replace("\"", "").replace("<", "").replace(">", "").replace("|", "")
        return " ".join(s.split())

    # 4. Construir lista de partes
    partes_nombre = []
    
    # Parte 1: Código Proyecto
    partes_nombre.append(limpiar_nombre(codigo_proy))
    
    # Parte 2: Código Cliente (Solo si existe)
    if codigo_cli:
        partes_nombre.append(limpiar_nombre(codigo_cli))
        
    # Parte 3: Empresa
    partes_nombre.append(limpiar_nombre(empresa))
    
    # Parte 4: Ubicación (DIRECCION-DISTRITO)
    partes_nombre.append(limpiar_nombre(ubicacion))
    
    # Unir con guiones bajos "_"
    nombre_base = "_".join(partes_nombre)
    
    print(f"   [INFO] Nombre Generado: {nombre_base}")
    
    # Rutas finales
    ruta_docx = os.path.join(ruta_proyecto, f"{nombre_base}.docx")
    ruta_pdf = os.path.join(ruta_proyecto, f"{nombre_base}.pdf")

    # MODO 1: SOLO CONVERTIR
    if modo == 'convert_only':
        # Buscamos si existe ya un DOCX con este nombre exacto
        if not os.path.exists(ruta_docx):
            # Intento de recuperación: buscar cualquier .docx en la carpeta que no sea temp_render
            posibles = [f for f in os.listdir(ruta_proyecto) if f.endswith(".docx") and f != "temp_render.docx" and not f.startswith("~$")]
            if posibles:
                ruta_docx = os.path.join(ruta_proyecto, posibles[0])
                nombre_base = os.path.splitext(posibles[0])[0]
                ruta_pdf = os.path.join(ruta_proyecto, f"{nombre_base}.pdf")
            else:
                print(f"Error: No se encuentra archivo Word para convertir.")
                sys.exit(1)
        
        try:
            print(f"Convirtiendo: {os.path.basename(ruta_docx)}")
            convert(ruta_docx, ruta_pdf)
            print(f"SUCCESS_PDF:{nombre_base}.pdf")
            sys.exit(0)
        except Exception as e:
            print(f"ERROR CONVERSION: {e}")
            sys.exit(1)

    # MODO 2: GENERACIÓN COMPLETA
    try:
        doc = DocxTemplate(ruta_plantilla)
        doc.render(datos)
        
        temp_docx = os.path.join(ruta_proyecto, "temp_render.docx")
        doc.save(temp_docx)
        
        doc_final = Document(temp_docx)
        limpiar_formato_variables(doc_final, datos)
        procesar_evidencias_universal(doc_final, ruta_proyecto, comentarios)
        
        # Eliminar words antiguos
        for f in os.listdir(ruta_proyecto):
            if f.endswith(".docx") and f != "temp_render.docx":
                try: os.remove(os.path.join(ruta_proyecto, f))
                except: pass

        doc_final.save(ruta_docx)
        if os.path.exists(temp_docx): os.remove(temp_docx)
        print(f"SUCCESS_WORD:{nombre_base}.docx")
        
    except Exception as e:
        print(f"ERROR PYTHON: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # Si el modo era 'pdf' (desde cero)
    if modo == 'pdf':
        try:
            print("Convirtiendo a PDF...")
            convert(ruta_docx, ruta_pdf)
            print(f"SUCCESS_PDF:{nombre_base}.pdf")
        except Exception as e:
            print(f"ERROR CONVERSION PDF: {e}")
            sys.exit(1)